import datetime, itertools
from collections import defaultdict
from docx import Document

def verificar_datas(d1, d2):
    data1 = datetime.datetime.strptime(d1, "%d/%m/%Y")
    data2 = datetime.datetime.strptime(d2, "%d/%m/%Y")
    
    data_atual = data1
    delta = datetime.timedelta(days=1)
    datas = []
    
    while data_atual <= data2:
        if data_atual.weekday() in [0, 2, 4]:  # Segunda (0), Quarta (2) e Sexta-feira (4)
            datas.append(data_atual)
        data_atual += delta
    
    print("Datas disponíveis:")
    for i, data in enumerate(datas):
        print(f"{i + 1}. {data.strftime('%d/%m/%Y')} ({data.strftime('%A')})")
    
    excluir = input("Deseja excluir alguma(s) data(s)? (S/N): ")
    
    if excluir.lower() == "s":
        indices_excluir = input("Digite os números das datas que deseja excluir (separados por vírgula): ")
        indices_excluir = [int(idx.strip()) for idx in indices_excluir.split(",")]
        
        datas_restantes = [data for i, data in enumerate(datas) if i + 1 not in indices_excluir]
        
        print("Datas restantes após exclusão:")
        for data in datas_restantes:
            print(f"{data.strftime('%d/%m/%Y')} ({data.strftime('%A')})")
        
        return datas_restantes
    
    return datas

def obter_cargos_disponiveis():
    print("Informe os cargos a serem preenchidos (separados por vírgula):")
    cargos = input().split(",")
    
    return [cargo.strip() for cargo in cargos]

def relacionar_datas_cargos(datas_resultantes, cargos_disponiveis):
    relacionamento = list(itertools.product(datas_resultantes, cargos_disponiveis))
    return relacionamento

def imprimir_relacionamento(relacionamento):
    for data, cargo in relacionamento:
        print(f"Data: {data.strftime('%d/%m/%Y')}, Cargo: {cargo}")

def verificar_funcionarios_disponiveis():
    funcionarios = {}
    
    # Lendo os nomes dos funcionários do arquivo "funcionarios.txt"
    try:
        with open("funcionarios.txt", "r", encoding='UTF-8') as arquivo:
            for indice, nome in enumerate(arquivo, start=1):
                funcionarios[indice] = nome.strip()
    except FileNotFoundError:
        print("Arquivo 'funcionarios.txt' não encontrado.")
        return []
    
    print("Lista de Funcionários:")
    for indice, nome in funcionarios.items():
        print(f"{indice}. {nome}")
    
    # Solicitando ao usuário os funcionários não disponíveis
    indisponiveis = input("Informe os números dos funcionários indisponíveis (separados por vírgula) ou 0: ")
    indisponiveis = [int(indice.strip()) for indice in indisponiveis.split(",")]
    
    # Criando a lista de funcionários disponíveis
    disponiveis = [nome for indice, nome in funcionarios.items() if indice not in indisponiveis]
    
    print("\nFuncionários Disponíveis:")
    n = 1
    for nome in disponiveis:
        print(str(n) + '. ' + nome)
        n +=1
    
    return disponiveis

def verificar_repeticao_semana(lista_tuplas):
    semana_funcionarios = defaultdict(list)

    for tupla in lista_tuplas:
        data, _, _, funcionario = tupla
        dia_semana = data.split('/')[1]  # Obtém o dia da semana

        # Considera apenas dias de segunda a sexta-feira
        if dia_semana in ['01', '02', '03', '04', '05']:
            semana = data.split('/')[0]  # Obtém o número da semana
            semana_funcionarios[semana].append(funcionario)

    for semana, funcionarios in semana_funcionarios.items():
        repetidos = set()
        for funcionario in funcionarios:
            if funcionarios.count(funcionario) > 1:
                repetidos.add(funcionario)

        if repetidos:
            repetidos_str = ', '.join(repetidos)
            print(f"Alerta: Os funcionários {repetidos_str} aparecem mais de uma vez na semana {semana}.")

def gerar_escala(relacionamento_data_cargos, funcionarios_disponiveis):
    
    # Verifica se há um último funcionário escalado
    try:
        with open('ultimo_funcionario.txt', 'r') as file:
            ultimo_funcionario = file.read().strip()
    except FileNotFoundError:
        ultimo_funcionario = funcionarios_disponiveis[0]

    # Encontra o índice do último funcionário escalado
    try:
        indice_ultimo_funcionario = funcionarios_disponiveis.index(ultimo_funcionario)
    except:
        funcionario_informado = int(input('Não encontrei correspondência. Indique o número do próximo funcionário a ser escalado: '))
        indice_ultimo_funcionario = funcionario_informado - 2
       
    # Gera a escala de funcionários
    escala = []
    for data, cargo in relacionamento_data_cargos:
        # Obtém o próximo funcionário disponível para escala
        proximo_indice = (indice_ultimo_funcionario + 1) % len(funcionarios_disponiveis)
        proximo_funcionario = funcionarios_disponiveis[proximo_indice]

        # Verifica se o próximo funcionário é igual ao último funcionário escalado
        while proximo_funcionario == ultimo_funcionario:
            proximo_indice = (proximo_indice + 1) % len(funcionarios_disponiveis)
            proximo_funcionario = funcionarios_disponiveis[proximo_indice]

        # Atualiza o índice do último funcionário escalado
        indice_ultimo_funcionario = proximo_indice

        # Armazena o último funcionário escalado
        with open('ultimo_funcionario.txt', 'w') as file:
            file.write(proximo_funcionario)

        # Formatação da data legível
        data_legivel = data.strftime('%d/%m/%Y')
        dia_semana = data.strftime('%A')

        # Adiciona a tupla à escala
        escala.append((data_legivel, dia_semana, cargo, proximo_funcionario))

    return escala



# Início do programa
data_inicial = input('Entre com a data inicial da escala: ')
data_final = input('Entre com a data final da escala: ')
datas_resultantes = verificar_datas(data_inicial, data_final)
cargos_disponiveis = obter_cargos_disponiveis()
relacionamento_data_cargos = relacionar_datas_cargos(datas_resultantes, cargos_disponiveis)
funcionarios_disponiveis = verificar_funcionarios_disponiveis()
escala = gerar_escala(relacionamento_data_cargos, funcionarios_disponiveis)

# Imprime a escala na tela e salva em docx

# Cria o objeto Document
doc = Document()

titulo = f'***** ESCALA DE SUBSTITUIÇÃO PARA {data_inicial} A {data_final} *****\n'
print('\n\n' + titulo)
paragrafo = doc.add_paragraph()
paragrafo.add_run(titulo)


for tupla in escala:
    data_legivel, dia_semana, cargo, funcionario = tupla

    # Adiciona as linhas ao documento
    paragrafo = doc.add_paragraph()
    paragrafo.add_run(f'Data: {data_legivel} ({dia_semana})')
    print(f'Data: {data_legivel} ({dia_semana})')
    paragrafo = doc.add_paragraph()
    paragrafo.add_run(f'Cargo: {cargo}')
    print(f'Cargo: {cargo}')
    paragrafo = doc.add_paragraph()
    paragrafo.add_run(f'Funcionário: {funcionario}')
    print(f'Funcionário: {funcionario}')
    paragrafo = doc.add_paragraph()
    paragrafo.add_run('---------------------------')
    print('---------------------------')

# Salva o documento
doc.save("..\escala.docx")
verificar_repeticao_semana(escala)
input()